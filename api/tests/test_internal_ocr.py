"""Integration test for the new internal OCR endpoint (Phase 2.2).

`POST /api/v1/internal/n8n/ocr` downloads a file from MinIO, runs OCR, and
returns the structured ocr_json. We patch the API key and MinIO download so the
test runs offline, but the real OCR + parsing logic executes.
"""
import io

import pytest
from docx import Document
from fastapi import FastAPI
from fastapi.testclient import TestClient

import routes.internal as internal
from services import minio as minio_svc


DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"


@pytest.fixture()
def client():
    app = FastAPI()
    app.include_router(internal.router)
    internal.API_SECRET = "test-internal-key"
    return TestClient(app)


def _make_docx_bytes():
    # Header lines followed by body lines — this is what lets
    # parse_text_to_json populate the `sections` dict.
    doc = Document()
    doc.add_paragraph("Experience:")
    doc.add_paragraph("Built microservices achieving 40% faster delivery.")
    doc.add_paragraph("Led a cross-functional team of engineers.")
    doc.add_paragraph("Skills: Python, Docker, Kubernetes, AWS.")
    doc.add_paragraph("Education: BS Computer Science.")
    doc.add_paragraph("Summary: Senior backend engineer.")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_ocr_endpoint_returns_structured_json(client, monkeypatch):
    monkeypatch.setattr(minio_svc, "download_file", lambda key: _make_docx_bytes())

    payload = {
        "resume_id": "resume-123",
        "minio_key": "resumes/test.docx",
        "file_type": DOCX_TYPE,
    }
    resp = client.post(
        "/api/v1/internal/n8n/ocr", json=payload,
        headers={"X-API-Key": "test-internal-key"},
    )
    assert resp.status_code == 200

    data = resp.json()
    assert data["status"] == "ocr_completed"
    assert data["resume_id"] == "resume-123"
    ocr_json = data["ocr_json"]
    # Real OCR parsing produced sections (not a stub).
    assert "sections" in ocr_json
    assert "experience" in ocr_json["sections"]
    assert "Built microservices" in ocr_json["sections"]["experience"]


def test_ocr_endpoint_rejects_wrong_api_key(client, monkeypatch):
    # A present-but-wrong key -> 403 (422 would mean the header was missing
    # entirely, which FastAPI rejects before the route runs).
    monkeypatch.setattr(minio_svc, "download_file", lambda key: _make_docx_bytes())
    resp = client.post(
        "/api/v1/internal/n8n/ocr",
        json={"resume_id": "r1", "minio_key": "k", "file_type": DOCX_TYPE},
        headers={"X-API-Key": "wrong-key"},
    )
    assert resp.status_code == 403


def test_ocr_endpoint_rejects_missing_minio_key(client, monkeypatch):
    monkeypatch.setattr(minio_svc, "download_file", lambda key: _make_docx_bytes())
    resp = client.post(
        "/api/v1/internal/n8n/ocr",
        json={"resume_id": "r1", "file_type": DOCX_TYPE},
        headers={"X-API-Key": "test-internal-key"},
    )
    assert resp.status_code == 400


def test_ocr_endpoint_handles_minio_download_failure(client, monkeypatch):
    def boom(key):
        raise RuntimeError("MinIO connection refused")

    monkeypatch.setattr(minio_svc, "download_file", boom)
    resp = client.post(
        "/api/v1/internal/n8n/ocr",
        json={"resume_id": "r1", "minio_key": "resumes/x.docx", "file_type": DOCX_TYPE},
        headers={"X-API-Key": "test-internal-key"},
    )
    assert resp.status_code == 500
